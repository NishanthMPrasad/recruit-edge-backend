# backend/document_generator.py
import io
import re
import base64
import os
from bs4 import BeautifulSoup
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Helper function to clean up extra whitespace ---
def clean_text(text: str) -> str:
    if not text:
        return ""
    cleaned_text = re.sub(r"[ \t]+", " ", text)
    cleaned_text = re.sub(r"\n{2,}", "\n", cleaned_text)
    cleaned_text = "\n".join(line.strip() for line in cleaned_text.split("\n"))
    return cleaned_text.strip()

# --- Helper to add HTML content to a DOCX paragraph ---
def add_html_to_docx_paragraph(paragraph, html_content, normal_style=None):
    if not html_content:
        return
    html_content = re.sub(r">\s+<", "><", html_content)
    html_content = re.sub(r"\s{2,}", " ", html_content)
    soup = BeautifulSoup(html_content, "html.parser")

    for element in (soup.body.contents if soup.body else soup.contents):
        if element.name == "ul":
            for li in element.find_all("li"):
                li_p = paragraph.insert_paragraph_before() if paragraph.text else paragraph
                li_p.style = normal_style
                li_p.add_run("•\t")
                li_p.paragraph_format.left_indent = Pt(36)
                li_p.paragraph_format.first_line_indent = Pt(-18)
                process_inline_html_for_docx(li_p, str(li.decode_contents()), normal_style)
        elif element.name == "p":
            p_tag = paragraph if not paragraph.text else paragraph._element.addnext(paragraph._element)._p.get_or_add_pPr()
            # Fallback if above line is too magical; create a new paragraph on doc level:
            if not hasattr(p_tag, "add_run"):
                p_tag = paragraph._parent.add_paragraph()
            p_tag.style = normal_style
            process_inline_html_for_docx(paragraph if p_tag is paragraph else p_tag, str(element.decode_contents()), normal_style)
        elif element.name == "br":
            paragraph.add_run().add_break()
        elif element.string:
            run = paragraph.add_run(str(element.string).strip())
            if normal_style:
                run.font.name = normal_style.font.name
                run.font.size = normal_style.font.size
        else:
            process_inline_html_for_docx(paragraph, str(element), normal_style)

# --- Helper for recursive inline HTML processing ---
def process_inline_html_for_docx(parent_obj, html_snippet, normal_style):
    temp_soup = BeautifulSoup(html_snippet, "html.parser")
    for content in temp_soup.contents:
        if content.name == "strong":
            run = parent_obj.add_run(content.get_text())
            run.bold = True
        elif content.name == "em":
            run = parent_obj.add_run(content.get_text())
            run.italic = True
        elif content.name == "a":
            run = parent_obj.add_run(content.get_text())
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        elif content.name == "br":
            parent_obj.add_run().add_break()
        elif content.string:
            parent_obj.add_run(str(content.string).strip())
        elif content.contents:
            process_inline_html_for_docx(parent_obj, str(content), normal_style)

# --- DOCX GENERATION ---
def generate_docx_from_data(data):
    doc = Document()

    # ----- Styles -----
    style = data.get("styleOptions", {})
    font_name = style.get("fontFamily", "Calibri").split(",")[0]
    font_size = style.get("fontSize", 11)
    accent_color_hex = style.get("accentColor", "#34495e").lstrip("#")
    accent_color_rgb = RGBColor.from_string(accent_color_hex)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)

    try:
        heading_style = doc.styles["SectionHeading"]
    except KeyError:
        heading_style = doc.styles.add_style("SectionHeading", 1)
    heading_style.font.name = font_name
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.font.color.rgb = accent_color_rgb
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)

    personal = data.get("personal", {})

    # ----- HEADER (table: logo | name+contact | profile) -----
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    table.columns[2].width = Inches(2.0)

    left_cell, center_cell, right_cell = table.rows[0].cells

    # Left: Pamten logo (optional)
    if data.get("pamtenLogoBase64"):
        img_data = data["pamtenLogoBase64"].split(",")[-1]
        img_bytes = base64.b64decode(img_data)
        p = left_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(1.4))

    # Center: Name + contact
    name_p = center_cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(personal.get("name", ""))
    r.bold = True
    r.font.name = font_name
    r.font.size = Pt(26)
    r.font.color.rgb = accent_color_rgb

    contact_items = [personal.get("email"), personal.get("phone"), personal.get("location")]
    if personal.get("legalStatus") and personal.get("legalStatus") != "Prefer not to say":
        contact_items.append(personal.get("legalStatus"))
    contact_p = center_cell.add_paragraph(" | ".join(filter(None, contact_items)))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Right: Profile picture (optional)
    if data.get("profilePicBase64"):
        img_data = data["profilePicBase64"].split(",")[-1]
        img_bytes = base64.b64decode(img_data)
        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(1.3))

    doc.add_paragraph()  # spacer

    # ----- Sections -----
    if data.get("summary"):
        doc.add_paragraph("Summary", style="SectionHeading")
        p = doc.add_paragraph(style=normal_style)
        add_html_to_docx_paragraph(p, data.get("summary"), normal_style)

    if data.get("experience") and any(e.get("jobTitle") for e in data.get("experience")):
        doc.add_paragraph("Experience", style="SectionHeading")
        for exp in data["experience"]:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(exp.get("jobTitle", "")).bold = True
            p.add_run(f"\n{exp.get('company', '')} | {exp.get('dates', '')}\n").italic = True
            add_html_to_docx_paragraph(p, exp.get("description", ""), normal_style)
            p.paragraph_format.space_after = Pt(12)

    if data.get("education") and any(e.get("degree") for e in data.get("education")):
        doc.add_paragraph("Education", style="SectionHeading")
        for edu in data["education"]:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(edu.get("degree", "")).bold = True
            p.add_run(f", {edu.get('institution', '')}\n")
            p.add_run(
                f"{edu.get('graduationYear', '')}"
                f"{f' | GPA: {edu.get('gpa')}' if edu.get('gpa') else ''}\n"
            ).italic = True
            add_html_to_docx_paragraph(p, edu.get("achievements", ""), normal_style)
            p.paragraph_format.space_after = Pt(12)

    if data.get("skills") and any(s.get("skills_list") for s in data.get("skills")):
        doc.add_paragraph("Skills", style="SectionHeading")
        for skill in data["skills"]:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(f"{skill.get('category', '')}: ").bold = True
            p.add_run(clean_text(skill.get("skills_list", "")))
            p.paragraph_format.space_after = Pt(6)

    if data.get("projects") and any(pr.get("title") for pr in data.get("projects")):
        doc.add_paragraph("Projects", style="SectionHeading")
        for proj in data["projects"]:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(proj.get("title", "")).bold = True
            p.add_run(f" ({proj.get('date', '')})\n").italic = True
            add_html_to_docx_paragraph(p, proj.get("description", ""), normal_style)
            p.paragraph_format.space_after = Pt(12)

    if data.get("publications") and any(pub.get("title") for pub in data.get("publications")):
        doc.add_paragraph("Publications", style="SectionHeading")
        for pub in data["publications"]:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(pub.get("title", "")).bold = True
            p.add_run(f" ({pub.get('date', '')})\n").italic = True
            small_sz = Pt(max(font_size - 1, 8))
            p.add_run(f"{pub.get('authors', '')} - {pub.get('journal', '')}\n").font.size = small_sz
            if pub.get("link"):
                p.add_run(f"Link: {pub.get('link')}").font.size = small_sz
            p.paragraph_format.space_after = Pt(12)

    if data.get("certifications") and any(c.get("name") for c in data.get("certifications")):
        doc.add_paragraph("Certifications", style="SectionHeading")
        for cert in data["certifications"]:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(cert.get("name", "")).bold = True
            issuer_date_text = cert.get("issuer", "") or ""
            if cert.get("date"):
                issuer_date_text += f" | {cert.get('date')}"
            p.add_run(f"\n{issuer_date_text}").italic = True
            p.paragraph_format.space_after = Pt(12)

    return doc

# --- PDF GENERATION ---
def generate_pdf_from_data(data):
    data["pamtenLogoSrc"] = data.get("pamtenLogoBase64") or None
    data["profilePicSrc"] = data.get("profilePicBase64") or None

    if data.get("skills"):
        for skill in data["skills"]:
            skill["skills_list"] = clean_text(skill.get("skills_list", ""))

    env = Environment(loader=FileSystemLoader(os.path.join(os.path.dirname(__file__), "assets")))
    template = env.get_template("resume_template.html")
    rendered_html = template.render(**data)
    return HTML(string=rendered_html).write_pdf()
